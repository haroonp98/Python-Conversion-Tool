import os
import tkinter as tk
from tkinter import filedialog, scrolledtext, ttk
import fitz
from googletrans import Translator
from docx import Document

def select_lang():
    source_lang = LANGUAGES[source_lang_combobox.get()]
    target_lang = LANGUAGES[target_lang_combobox.get()]
    return source_lang, target_lang

def translate_text(text, target_lang):
    translator = Translator()
    translation = translator.translate(text, dest=target_lang)
    return translation.text

def open_file(file_path):
    if file_path.endswith('.pdf'):
        doc = fitz.open(file_path)
        full_text = ""
        for page_number in range(doc.page_count):
            page = doc[page_number]
            full_text += page.get_text("text")
        return full_text
    elif file_path.endswith('.docx'):
        return open_docx(file_path)
    else:
        return "UNSUPPORTED FILE FORMAT"

def open_docx(file_path):
    doc = Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n'.join(full_text)

def write_to_text_widget(translated_text):
    result_text.config(state=tk.NORMAL)
    result_text.delete(1.0, tk.END)
    result_text.insert(tk.END, translated_text)
    result_text.config(state=tk.DISABLED)

def convert():
    source_lang, target_lang = select_lang()

    if conversion_type_var.get() == "text":
        text_to_translate = text_input.get("1.0", tk.END)
        translated_text = translate_text(text_to_translate, target_lang)
        write_to_text_widget(translated_text)
        result_label.config(text="Conversion completed successfully!")

    elif conversion_type_var.get() == "files":
        input_files = get_selected_files()
        output_folder = output_folder_var.get()

        for input_file in input_files:
            file_text = open_file(input_file)
            CHUNK_LIMIT = 500
            translated_file_text = ""

            chunks_list = [file_text[i:i + CHUNK_LIMIT] for i in range(0, len(file_text), CHUNK_LIMIT)]

            for current_chunk in chunks_list:
                translated_chunk = translate_text(current_chunk, target_lang)
                translated_file_text += translated_chunk

            output_file_path = os.path.join(output_folder, os.path.basename(input_file)[:-4] + "_translated")

            if input_file.endswith('.pdf'):
                output_file_path += '.pdf'
                save_to_pdf(translated_file_text, output_file_path)
            elif input_file.endswith('.docx'):
                output_file_path += '.docx'
                save_to_docx(translated_file_text, output_file_path)

            result_label.config(text=f"Conversion of {os.path.basename(input_file)} completed successfully!")

            # Display translated text in the GUI
            write_to_text_widget(translated_file_text)

def save_to_pdf(translated_text, output_file_path):
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((10, 10), translated_text, fontsize=10)

    # Save the document to a PDF file
    doc.save(output_file_path)
    doc.close()

def save_to_docx(translated_text, output_file_path):
    doc = Document()
    doc.add_paragraph(translated_text)

    # Save the document to a DOCX file
    doc.save(output_file_path)

# GUI Setup
app = tk.Tk()
app.title("Text Translator")

# Language dictionary
LANGUAGES = {
    'Afrikaans': 'af',
    'Albanian': 'sq',
    'Amharic': 'am',
    'Arabic': 'ar',
    'Armenian': 'hy',
    'Azerbaijani': 'az',
    'Basque': 'eu',
    'Belarusian': 'be',
    'Bengali': 'bn',
    'Bosnian': 'bs',
    'Bulgarian': 'bg',
    'Catalan': 'ca',
    'Cebuano': 'ceb',
    'Chichewa': 'ny',
    'Chinese (Simplified)': 'zh-cn',
    'Chinese (Traditional)': 'zh-tw',
    'Corsican': 'co',
    'Croatian': 'hr',
    'Czech': 'cs',
    'Danish': 'da',
    'Dutch': 'nl',
    'English': 'en',
    'Esperanto': 'eo',
    'Estonian': 'et',
    'Filipino': 'tl',
    'Finnish': 'fi',
    'French': 'fr',
    'Frisian': 'fy',
    'Galician': 'gl',
    'Georgian': 'ka',
    'German': 'de',
    'Greek': 'el',
    'Gujarati': 'gu',
    'Haitian Creole': 'ht',
    'Hausa': 'ha',
    'Hawaiian': 'haw',
    'Hebrew': 'iw',
    'Hindi': 'hi',
    'Hmong': 'hmn',
    'Hungarian': 'hu',
    'Icelandic': 'is',
    'Igbo': 'ig',
    'Indonesian': 'id',
    'Irish': 'ga',
    'Italian': 'it',
    'Japanese': 'ja',
    'Javanese': 'jw',
    'Kannada': 'kn',
    'Kazakh': 'kk',
    'Khmer': 'km',
    'Korean': 'ko',
    'Kurdish (Kurmanji)': 'ku',
    'Kyrgyz': 'ky',
    'Lao': 'lo',
    'Latin': 'la',
    'Latvian': 'lv',
    'Lithuanian': 'lt',
    'Luxembourgish': 'lb',
    'Macedonian': 'mk',
    'Malagasy': 'mg',
    'Malay': 'ms',
    'Malayalam': 'ml',
    'Maltese': 'mt',
    'Maori': 'mi',
    'Marathi': 'mr',
    'Mongolian': 'mn',
    'Myanmar (Burmese)': 'my',
    'Nepali': 'ne',
    'Norwegian': 'no',
    'Odia': 'or',
    'Pashto': 'ps',
    'Persian': 'fa',
    'Polish': 'pl',
    'Portuguese': 'pt',
    'Punjabi': 'pa',
    'Romanian': 'ro',
    'Russian': 'ru',
    'Samoan': 'sm',
    'Scots Gaelic': 'gd',
    'Serbian': 'sr',
    'Sesotho': 'st',
    'Shona': 'sn',
    'Sindhi': 'sd',
    'Sinhala': 'si',
    'Slovak': 'sk',
    'Slovenian': 'sl',
    'Somali': 'so',
    'Spanish': 'es',
    'Sundanese': 'su',
    'Swahili': 'sw',
    'Swedish': 'sv',
    'Tajik': 'tg',
    'Tamil': 'ta',
    'Telugu': 'te',
    'Thai': 'th',
    'Turkish': 'tr',
    'Ukrainian': 'uk',
    'Urdu': 'ur',
    'Uyghur': 'ug',
    'Uzbek': 'uz',
    'Vietnamese': 'vi',
    'Welsh': 'cy',
    'Xhosa': 'xh',
    'Yiddish': 'yi',
    'Yoruba': 'yo',
    'Zulu': 'zu'
}

# Variables
conversion_type_var = tk.StringVar(value="text")
source_lang_combobox = ttk.Combobox(app, values=list(LANGUAGES.keys()), state="readonly")
target_lang_combobox = ttk.Combobox(app, values=list(LANGUAGES.keys()), state="readonly")

# GUI Components
source_lang_label = tk.Label(app, text="Source Language:")
source_lang_combobox = ttk.Combobox(app, values=list(LANGUAGES.keys()), state="readonly")
source_lang_combobox.set("English")

target_lang_label = tk.Label(app, text="Target Language:")
target_lang_combobox = ttk.Combobox(app, values=list(LANGUAGES.keys()), state="readonly")
target_lang_combobox.set("Spanish")

conversion_type_label = tk.Label(app, text="Conversion Type:")
conversion_type_text_radiobutton = tk.Radiobutton(app, text="Text", variable=conversion_type_var, value="text")
conversion_type_files_radiobutton = tk.Radiobutton(app, text="Files", variable=conversion_type_var, value="files")

input_files_label = tk.Label(app, text="Select Input Files:")
input_files_listbox = tk.Listbox(app, width=50, height=5, selectmode=tk.MULTIPLE)
input_files_scrollbar = tk.Scrollbar(app, orient="vertical", command=input_files_listbox.yview)
input_files_listbox.config(yscrollcommand=input_files_scrollbar.set)
input_files_button = tk.Button(app, text="Browse Files", command=lambda: browse_files("input"))

input_text_label = tk.Label(app, text="Enter Text:")
text_input = scrolledtext.ScrolledText(app, wrap=tk.WORD, width=50, height=5)

convert_button = tk.Button(app, text="Convert", command=convert)

result_label = tk.Label(app, text="")

result_text = scrolledtext.ScrolledText(app, wrap=tk.WORD, width=50, height=10, state=tk.DISABLED)

output_folder_label = tk.Label(app, text="Select Output Folder:")
output_folder_var = tk.StringVar()
output_folder_entry = tk.Entry(app, textvariable=output_folder_var)
output_folder_button = tk.Button(app, text="Browse Folder", command=lambda: browse_folder("output"))

# GUI Layout
source_lang_label.grid(row=0, column=0, padx=10, pady=5, sticky="e")
source_lang_combobox.grid(row=0, column=1, padx=10, pady=5)

target_lang_label.grid(row=1, column=0, padx=10, pady=5, sticky="e")
target_lang_combobox.grid(row=1, column=1, padx=10, pady=5)

conversion_type_label.grid(row=2, column=0, padx=10, pady=5, sticky="e")
conversion_type_text_radiobutton.grid(row=2, column=1, padx=10, pady=5)
conversion_type_files_radiobutton.grid(row=2, column=2, padx=10, pady=5)

input_files_label.grid(row=3, column=0, padx=10, pady=5, sticky="e")
input_files_listbox.grid(row=3, column=1, padx=10, pady=5, columnspan=2, sticky="ew")
input_files_scrollbar.grid(row=3, column=3, pady=5, sticky="ns")
input_files_button.grid(row=3, column=4, padx=5, pady=5)

input_text_label.grid(row=4, column=0, padx=10, pady=5, sticky="e")
text_input.grid(row=4, column=1, padx=10, pady=5, columnspan=2, sticky="ew")

convert_button.grid(row=5, column=1, pady=10)

result_label.grid(row=6, column=1)

result_text.grid(row=7, column=1, padx=10, pady=5, columnspan=2, sticky="ew")

output_folder_label.grid(row=8, column=0, padx=10, pady=5, sticky="e")
output_folder_entry.grid(row=8, column=1, padx=10, pady=5)
output_folder_button.grid(row=8, column=2, padx=5, pady=5)




def browse_files(folder_type):
    selected_files = filedialog.askopenfilenames(title=f"Select {folder_type} Files", filetypes=[("PDF files", "*.pdf"), ("Word Doc", "*.docx")])
    input_files_listbox.delete(0, tk.END)
    for file in selected_files:
        input_files_listbox.insert(tk.END, file)

def browse_folder(folder_type):
    selected_folder = filedialog.askdirectory(title=f"Select {folder_type} Folder")
    if folder_type == "output":
        output_folder_var.set(selected_folder)

def get_selected_files():
    return input_files_listbox.get(0, tk.END)

app.mainloop()
